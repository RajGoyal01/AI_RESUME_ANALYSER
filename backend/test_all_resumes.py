"""
Comprehensive Resume Analyser Test Script
Tests the backend with multiple resumes of different file types and domains.
"""
import os
import sys
import json
import time

# ─── Test Configuration ───
TEST_DIR = os.path.join(os.path.dirname(__file__), "test_resumes")

# Direct backend tests (no server needed)
sys.path.insert(0, os.path.dirname(__file__))
from resume_parser import parse_resume, extract_text
from analyser import analyze_resume

# Color codes for terminal
class C:
    GREEN = "\033[92m"
    RED = "\033[91m"
    YELLOW = "\033[93m"
    CYAN = "\033[96m"
    MAGENTA = "\033[95m"
    BOLD = "\033[1m"
    END = "\033[0m"

def separator(title):
    print(f"\n{C.CYAN}{'═'*70}")
    print(f"  {C.BOLD}{title}{C.END}")
    print(f"{C.CYAN}{'═'*70}{C.END}")

def sub_sep(title):
    print(f"\n  {C.MAGENTA}── {title} ──{C.END}")

def test_resume(filepath, domain, course, branch, label):
    """Test a single resume end-to-end and return results."""
    separator(f"TEST: {label}")
    print(f"  📁 File: {os.path.basename(filepath)}")
    print(f"  📦 Size: {os.path.getsize(filepath)} bytes")
    print(f"  🏷️  Domain: {domain} | Course: {course} | Branch: {branch}")

    errors = []

    # ─── Step 1: Parse Resume ───
    sub_sep("Step 1: Parsing Resume")
    try:
        parsed = parse_resume(filepath)
        if "error" in parsed:
            print(f"  {C.RED}❌ Parse Error: {parsed['error']}{C.END}")
            errors.append(f"Parse failed: {parsed['error']}")
            return {"label": label, "status": "FAIL", "errors": errors}

        print(f"  {C.GREEN}✅ Text extracted: {parsed.get('word_count', 0)} words{C.END}")
        print(f"  👤 Name: {parsed.get('name', 'Not found')}")
        print(f"  📧 Email: {parsed.get('email', 'Not found')}")
        print(f"  📱 Phone: {parsed.get('phone', 'Not found')}")
        print(f"  🔗 Links: {json.dumps(parsed.get('links', {}))}")
        print(f"  📚 Sections: {', '.join(parsed.get('section_names', []))}")
        print(f"  🎓 Education: {len(parsed.get('education', []))} entries")
        print(f"  🔧 Projects: {parsed.get('project_count', 0)}")

        # Validate essential fields
        if not parsed.get('name'):
            errors.append("Name not extracted")
            print(f"  {C.YELLOW}⚠️ Name not extracted{C.END}")
        if not parsed.get('email'):
            errors.append("Email not extracted")
            print(f"  {C.YELLOW}⚠️ Email not extracted{C.END}")
        if parsed.get('word_count', 0) < 30:
            errors.append("Very low word count")
            print(f"  {C.YELLOW}⚠️ Very low word count{C.END}")
    except Exception as e:
        print(f"  {C.RED}❌ Parse Exception: {e}{C.END}")
        errors.append(f"Parse exception: {e}")
        return {"label": label, "status": "FAIL", "errors": errors}

    # ─── Step 2: Analyze Resume ───
    sub_sep("Step 2: Analyzing Resume")
    try:
        results = analyze_resume(parsed, domain, course, branch)

        overall = results.get("overall_score", 0)
        grade = results.get("grade", "?")
        scores = results.get("scores", {})
        
        # Color based on score
        if overall >= 70:
            color = C.GREEN
        elif overall >= 50:
            color = C.YELLOW
        else:
            color = C.RED

        print(f"  {color}{C.BOLD}📊 OVERALL SCORE: {overall}/100 (Grade: {grade}){C.END}")
        print()
        print(f"  {'Category':<25} {'Score':>6}")
        print(f"  {'─'*35}")
        for cat, score in scores.items():
            bar = '█' * (score // 5) + '░' * (20 - score // 5)
            cat_color = C.GREEN if score >= 70 else C.YELLOW if score >= 40 else C.RED
            print(f"  {cat:<25} {cat_color}{score:>5}/100{C.END} {bar}")

        # MNC Scores
        sub_sep("MNC Readiness Scores (Top 10)")
        mnc = results.get("mnc_scores", {})
        sorted_mnc = sorted(mnc.items(), key=lambda x: -x[1])[:10]
        for company, score in sorted_mnc:
            bar = '█' * (score // 5) + '░' * (20 - score // 5)
            mnc_color = C.GREEN if score >= 60 else C.YELLOW if score >= 40 else C.RED
            print(f"  {company:<20} {mnc_color}{score:>5}%{C.END} {bar}")

        # Company Recommendations (Top 5)
        sub_sep("Top 5 Company Recommendations")
        recs = results.get("company_recommendations", [])[:5]
        for r in recs:
            dm = "✓ Domain" if r.get("domain_match") else "✗ Domain"
            print(f"  {r.get('emoji','')} {r['company']:<18} {r['chance_score']:>3}% ({r['level']:<8}) | {dm} | Bar: {r['hiring_bar']}")

        # Weak Areas
        sub_sep("Weak Areas")
        for w in results.get("weak_areas", []):
            sev_color = C.RED if w['severity'] == 'high' else C.YELLOW if w['severity'] == 'medium' else C.CYAN
            print(f"  {sev_color}[{w['severity'].upper()}]{C.END} {w['area']}")
            print(f"       💡 {w['suggestion'][:100]}...")

        # Strong Areas
        sub_sep("Strong Areas")
        for s in results.get("strong_areas", []):
            print(f"  {C.GREEN}✅ {s}{C.END}")

        # Suggestions
        sub_sep("Suggestions")
        for s in results.get("suggestions", []):
            print(f"  💡 {s}")

        # Validate scoring logic
        if overall < 0 or overall > 100:
            errors.append(f"Invalid overall score: {overall}")
        for cat, score in scores.items():
            if score < 0 or score > 100:
                errors.append(f"Invalid {cat} score: {score}")
        if not results.get("mnc_scores"):
            errors.append("No MNC scores generated")
        if not results.get("company_recommendations"):
            errors.append("No company recommendations generated")

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"  {C.RED}❌ Analysis Exception: {e}{C.END}")
        errors.append(f"Analysis exception: {e}")
        return {"label": label, "status": "FAIL", "errors": errors}

    status = "PASS" if not errors else "WARN"
    return {
        "label": label, "status": status, "errors": errors,
        "overall_score": overall, "grade": grade, "scores": scores,
        "word_count": parsed.get("word_count", 0),
    }


def main():
    print(f"\n{C.BOLD}{C.CYAN}")
    print("╔══════════════════════════════════════════════════════════════╗")
    print("║     AI RESUME ANALYSER - COMPREHENSIVE TEST SUITE          ║")
    print("║     Testing Multiple Resumes, File Types & Domains         ║")
    print("╚══════════════════════════════════════════════════════════════╝")
    print(C.END)

    all_results = []

    # ─── TEST CASES ───
    test_cases = [
        {
            "file": os.path.join(TEST_DIR, "resume_cs_senior.txt"),
            "domain": "Computer Science", "course": "B.Tech", "branch": "Computer Science",
            "label": "CS Senior Engineer (TXT) - Strong Profile"
        },
        {
            "file": os.path.join(TEST_DIR, "resume_fresher_weak.txt"),
            "domain": "Computer Science", "course": "B.Tech", "branch": "Computer Science",
            "label": "CS Fresher (TXT) - Weak Profile"
        },
        {
            "file": os.path.join(TEST_DIR, "resume_data_science.txt"),
            "domain": "Data Science", "course": "M.Tech", "branch": "Data Science",
            "label": "Data Scientist (TXT) - Strong Profile"
        },
        {
            "file": os.path.join(TEST_DIR, "resume_mechanical.txt"),
            "domain": "Mechanical Engineering", "course": "B.Tech", "branch": "Mechanical Engineering",
            "label": "Mechanical Engineer (TXT) - Mid Profile"
        },
        {
            "file": os.path.join(TEST_DIR, "resume_mba.txt"),
            "domain": "MBA / Business", "course": "MBA", "branch": "MBA / Business",
            "label": "MBA Marketing (TXT) - Strong Profile"
        },
        {
            "file": os.path.join(os.path.dirname(__file__), "test_resume.txt"),
            "domain": "Computer Science", "course": "B.Tech", "branch": "Computer Science",
            "label": "Original Test Resume (TXT) - Reference"
        },
    ]

    # Also try DOCX if python-docx is available
    try:
        from docx import Document
        # Create a DOCX test resume
        docx_path = os.path.join(TEST_DIR, "resume_ece.docx")
        doc = Document()
        doc.add_heading('RAHUL KRISHNAN', 0)
        doc.add_paragraph('Email: rahul.k@gmail.com | Phone: +91 9876501234')
        doc.add_paragraph('LinkedIn: https://www.linkedin.com/in/rahulk-ece | GitHub: https://www.github.com/rahulkece')
        
        doc.add_heading('EDUCATION', level=1)
        doc.add_paragraph('B.Tech in Electronics & Communication Engineering - NIT Warangal (2020-2024)\nCGPA: 8.7/10')
        
        doc.add_heading('SKILLS', level=1)
        doc.add_paragraph('Programming: C, C++, Python, MATLAB, Verilog, VHDL, Embedded C')
        doc.add_paragraph('Tools: Arduino, Raspberry Pi, Cadence, Xilinx, Keil, STM32, ESP32')
        doc.add_paragraph('Core: Analog Electronics, Digital Electronics, Signals and Systems, Communication Systems, Digital Signal Processing, Microprocessors, Microcontrollers, VLSI Design, Electromagnetic Theory, Control Systems')
        doc.add_paragraph('Technologies: PCB Design, Embedded Systems, IoT, Wireless Communication, 5G, FPGA, Signal Processing, Image Processing')
        
        doc.add_heading('EXPERIENCE', level=1)
        doc.add_paragraph('Embedded Systems Intern - Bosch India (May 2023 - Aug 2023)')
        doc.add_paragraph('- Developed firmware for automotive ECU using C and ARM Cortex-M4')
        doc.add_paragraph('- Designed PCB layout for sensor interface module using Altium Designer')
        doc.add_paragraph('- Implemented CAN bus communication protocol improving data transfer by 30%')
        doc.add_paragraph('- Built automated testing framework reducing testing time by 50%')
        
        doc.add_heading('PROJECTS', level=1)
        doc.add_paragraph('1. IoT-Based Smart Home System')
        doc.add_paragraph('   - Designed and built using ESP32, MQTT protocol, and custom PCB')
        doc.add_paragraph('   - Implemented voice control using Python and NLP')
        doc.add_paragraph('   - Won first place in university hackathon')
        doc.add_paragraph('2. FPGA-Based Image Processing')
        doc.add_paragraph('   - Implemented real-time edge detection on Xilinx Zynq FPGA')
        doc.add_paragraph('   - Achieved 60 FPS processing with custom Verilog IP cores')
        doc.add_paragraph('3. Autonomous Drone Navigation')
        doc.add_paragraph('   - Built using Raspberry Pi and custom sensor fusion algorithm')
        doc.add_paragraph('   - Programmed autonomous flight controller in Python and C++')
        
        doc.add_heading('CERTIFICATIONS', level=1)
        doc.add_paragraph('- Certified LabVIEW Associate Developer')
        doc.add_paragraph('- Coursera IoT Specialization')
        
        doc.add_heading('ACHIEVEMENTS', level=1)
        doc.add_paragraph('- Winner of National Robotics Competition 2023')
        doc.add_paragraph('- Published paper on "5G Antenna Design" at IEEE Conference')
        doc.add_paragraph('- Rank 2 in ECE department')
        
        doc.add_heading('SOFT SKILLS', level=1)
        doc.add_paragraph('Problem Solving, Teamwork, Communication, Analytical, Leadership, Creativity')
        
        doc.save(docx_path)
        print(f"  {C.GREEN}✅ Created DOCX test resume{C.END}")
        
        test_cases.append({
            "file": docx_path,
            "domain": "Electronics & Communication", "course": "B.Tech", "branch": "Electronics & Communication",
            "label": "ECE Engineer (DOCX) - Mid-Strong Profile"
        })
    except ImportError:
        print(f"  {C.YELLOW}⚠️ python-docx not installed, skipping DOCX test{C.END}")

    # Also create a PDF test if possible
    try:
        # Try creating a simple PDF using reportlab or fpdf
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=16)
            pdf.cell(0, 10, "MEERA NAIR", ln=True, align="C")
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 8, "Email: meera.nair@gmail.com | Phone: +91 7788990011", ln=True, align="C")
            pdf.cell(0, 8, "LinkedIn: https://www.linkedin.com/in/meeranair", ln=True, align="C")
            pdf.ln(5)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "EDUCATION", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 6, "B.Tech in Electrical Engineering - IIT Madras (2020-2024) | CGPA: 8.9/10", ln=True)
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "SKILLS", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 6, "Core: Circuit Analysis, Analog Electronics, Digital Electronics, Power Systems, Control Systems", ln=True)
            pdf.cell(0, 6, "Programming: MATLAB, Simulink, Python, C, C++, Verilog, LabVIEW", ln=True)
            pdf.cell(0, 6, "Tools: Arduino, Raspberry Pi, Cadence, LTSpice, Proteus, Altium Designer", ln=True)
            pdf.cell(0, 6, "Technologies: PCB Design, Embedded Systems, IoT, PLC, SCADA, Renewable Energy, Automation", ln=True)
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "EXPERIENCE", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 6, "Power Systems Intern - Siemens India (May 2023 - Aug 2023)", ln=True)
            pdf.cell(0, 6, "- Designed power distribution system for a 50MW solar plant", ln=True)
            pdf.cell(0, 6, "- Automated monitoring using SCADA reducing downtime by 25%", ln=True)
            pdf.cell(0, 6, "- Implemented PLC programming for industrial automation", ln=True)
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "PROJECTS", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 6, "1. Smart Grid Energy Management System using IoT and Machine Learning", ln=True)
            pdf.cell(0, 6, "2. Solar MPPT Controller designed with MATLAB and Arduino", ln=True)
            pdf.cell(0, 6, "3. Home Automation System using Raspberry Pi and sensor technology", ln=True)
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "CERTIFICATIONS", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 6, "- Certified LabVIEW Associate Developer", ln=True)
            pdf.cell(0, 6, "- Coursera Power Electronics Specialization", ln=True)
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "ACHIEVEMENTS", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 6, "- Winner of Smart India Hackathon 2023", ln=True)
            pdf.cell(0, 6, "- Published paper on Renewable Energy Systems at IEEE", ln=True)

            pdf_path = os.path.join(TEST_DIR, "resume_electrical.pdf")
            pdf.output(pdf_path)
            print(f"  {C.GREEN}✅ Created PDF test resume{C.END}")
            
            test_cases.append({
                "file": pdf_path,
                "domain": "Electrical Engineering", "course": "B.Tech", "branch": "Electrical Engineering",
                "label": "Electrical Engineer (PDF) - Mid Profile"
            })
        except ImportError:
            # If fpdf not available, try creating a minimal valid PDF manually
            pdf_path = os.path.join(TEST_DIR, "resume_electrical.pdf")
            # We'll skip PDF creation if no library available
            print(f"  {C.YELLOW}⚠️ fpdf not installed. Trying PyPDF2 method...{C.END}")
            # Create a minimal PDF with PyPDF2 (write-only)
            try:
                import PyPDF2
                from io import BytesIO
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas
                buf = BytesIO()
                c = canvas.Canvas(buf, pagesize=letter)
                c.drawString(72, 750, "MEERA NAIR")
                c.drawString(72, 735, "Email: meera.nair@gmail.com | Phone: +91 7788990011")
                c.drawString(72, 720, "LinkedIn: https://www.linkedin.com/in/meeranair")
                c.drawString(72, 695, "EDUCATION")
                c.drawString(72, 680, "B.Tech Electrical Engineering - IIT Madras | CGPA: 8.9/10")
                c.drawString(72, 655, "SKILLS")
                c.drawString(72, 640, "Core: Circuit Analysis, Power Systems, Control Systems, Electromagnetic Theory")
                c.drawString(72, 625, "Programming: MATLAB, Simulink, Python, C, Verilog, LabVIEW")
                c.drawString(72, 610, "Technologies: PCB Design, Embedded Systems, IoT, PLC, SCADA, Automation")
                c.drawString(72, 585, "EXPERIENCE")
                c.drawString(72, 570, "Power Systems Intern - Siemens India (2023)")
                c.drawString(72, 555, "- Designed power distribution system for 50MW solar plant")
                c.drawString(72, 540, "- Automated monitoring using SCADA reducing downtime by 25%")
                c.drawString(72, 515, "PROJECTS")
                c.drawString(72, 500, "1. Smart Grid Energy Management System using IoT")
                c.drawString(72, 485, "2. Solar MPPT Controller using MATLAB and Arduino")
                c.drawString(72, 470, "3. Home Automation System using Raspberry Pi")
                c.drawString(72, 445, "CERTIFICATIONS")
                c.drawString(72, 430, "- Coursera Power Electronics Specialization")
                c.drawString(72, 405, "ACHIEVEMENTS")
                c.drawString(72, 390, "- Winner of Smart India Hackathon 2023")
                c.save()
                with open(pdf_path, "wb") as f:
                    f.write(buf.getvalue())
                print(f"  {C.GREEN}✅ Created PDF test resume via reportlab{C.END}")
                test_cases.append({
                    "file": pdf_path,
                    "domain": "Electrical Engineering", "course": "B.Tech", "branch": "Electrical Engineering",
                    "label": "Electrical Engineer (PDF) - Mid Profile"
                })
            except ImportError:
                print(f"  {C.YELLOW}⚠️ No PDF creation library available, skipping PDF test{C.END}")
    except Exception as e:
        print(f"  {C.YELLOW}⚠️ PDF creation failed: {e}{C.END}")

    # ─── RUN ALL TESTS ───
    print(f"\n  {C.BOLD}Running {len(test_cases)} test cases...{C.END}\n")

    for tc in test_cases:
        if not os.path.exists(tc["file"]):
            print(f"  {C.RED}❌ File not found: {tc['file']}{C.END}")
            all_results.append({"label": tc["label"], "status": "FAIL", "errors": ["File not found"]})
            continue
        result = test_resume(tc["file"], tc["domain"], tc["course"], tc["branch"], tc["label"])
        all_results.append(result)

    # ─── SUMMARY ───
    separator("📋 TEST SUMMARY")
    print(f"\n  {'Test Case':<50} {'Status':>8} {'Score':>7}")
    print(f"  {'─'*70}")
    
    pass_count = 0
    warn_count = 0
    fail_count = 0
    
    for r in all_results:
        status = r["status"]
        score_str = f"{r.get('overall_score', '-')}/100" if 'overall_score' in r else "  -"
        
        if status == "PASS":
            color = C.GREEN
            symbol = "✅"
            pass_count += 1
        elif status == "WARN":
            color = C.YELLOW
            symbol = "⚠️"
            warn_count += 1
        else:
            color = C.RED
            symbol = "❌"
            fail_count += 1
        
        print(f"  {symbol} {r['label']:<48} {color}{status:>6}{C.END} {score_str:>7}")
        if r.get("errors"):
            for e in r["errors"]:
                print(f"     {C.RED}└─ {e}{C.END}")

    print(f"\n  {C.BOLD}Results: {C.GREEN}{pass_count} PASS{C.END} | {C.YELLOW}{warn_count} WARN{C.END} | {C.RED}{fail_count} FAIL{C.END}")
    print(f"  {C.BOLD}Total Tests: {len(all_results)}{C.END}\n")

    # ─── SCORING VALIDATION ───
    separator("📈 SCORE COMPARISON & VALIDATION")
    
    scores_data = [(r["label"], r.get("overall_score", 0)) for r in all_results if "overall_score" in r]
    scores_data.sort(key=lambda x: -x[1])
    
    print(f"\n  {'Resume':<50} {'Score':>7} {'Visual':>5}")
    print(f"  {'─'*70}")
    for label, score in scores_data:
        bar = '█' * (score // 2) + '░' * (50 - score // 2)
        color = C.GREEN if score >= 70 else C.YELLOW if score >= 40 else C.RED
        print(f"  {label:<50} {color}{score:>5}{C.END}  {bar}")
    
    # Validate relative ordering makes sense
    print(f"\n  {C.BOLD}Score Ordering Validation:{C.END}")
    
    # Strong CS senior should score higher than weak fresher
    cs_senior = next((r for r in all_results if "CS Senior" in r.get("label","")), None)
    cs_weak = next((r for r in all_results if "Fresher" in r.get("label","")), None)
    
    if cs_senior and cs_weak and "overall_score" in cs_senior and "overall_score" in cs_weak:
        if cs_senior["overall_score"] > cs_weak["overall_score"]:
            print(f"  {C.GREEN}✅ CS Senior ({cs_senior['overall_score']}) > CS Fresher ({cs_weak['overall_score']}) - CORRECT{C.END}")
        else:
            print(f"  {C.RED}❌ CS Senior ({cs_senior['overall_score']}) <= CS Fresher ({cs_weak['overall_score']}) - INCORRECT{C.END}")

    print(f"\n{'═'*70}")
    print(f"{C.BOLD}{C.GREEN}  ✅ All tests completed!{C.END}")
    print(f"{'═'*70}\n")


if __name__ == "__main__":
    main()
